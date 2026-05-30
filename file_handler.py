""""文件处理模块"""
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class FileHandler:
    type_map = ['.docx', '.md', '.txt']

    def _parse_numbering_formats(self, doc):
        """解析 Word 文档中的自动编号定义，返回 {numId: {ilvl: format_info}}"""
        try:
            numbering_part = doc.part.numbering_part
            if numbering_part is None:
                return {}
            root = numbering_part._element

            # 第一步：解析 abstractNum（编号格式模板）
            abstract_nums = {}
            for abs_num in root.findall(qn('w:abstractNum')):
                abs_id = abs_num.get(qn('w:abstractNumId'))
                levels = {}
                for lvl in abs_num.findall(qn('w:lvl')):
                    ilvl = lvl.get(qn('w:ilvl'))
                    fmt_el = lvl.find(qn('w:numFmt'))
                    fmt = fmt_el.get(qn('w:val')) if fmt_el is not None else 'decimal'
                    lvl_text_el = lvl.find(qn('w:lvlText'))
                    lvl_text = lvl_text_el.get(qn('w:val')) if lvl_text_el is not None else '%1.'
                    levels[ilvl] = {'fmt': fmt, 'lvlText': lvl_text}
                abstract_nums[abs_id] = levels

            # 第二步：解析 num → abstractNum 的映射
            numbering_map = {}
            for num in root.findall(qn('w:num')):
                num_id = num.get(qn('w:numId'))
                abs_ref = num.find(qn('w:abstractNumId'))
                if abs_ref is not None:
                    abs_id = abs_ref.get(qn('w:val'))
                    if abs_id in abstract_nums:
                        numbering_map[num_id] = abstract_nums[abs_id]

            return numbering_map
        except Exception:
            return {}

    def _format_list_number(self, fmt, lvl_text, counter):
        """根据编号格式和计数器生成编号文本"""
        # 项目符号（bullet）不递增
        if fmt == 'bullet':
            return lvl_text if lvl_text else '• '

        # 中文数字（chineseCounting / chineseCountingThousand）
        cn_nums = ['零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        if fmt in ('chineseCounting', 'chineseCountingThousand', 'chineseLegalSimplified'):
            n = cn_nums[counter] if counter < len(cn_nums) else str(counter)
            if lvl_text and '%' in lvl_text:
                return lvl_text.replace('%1', n).replace('%2', n)
            return f'{n}、'

        # 字母编号
        if fmt == 'upperLetter':
            n = chr(64 + counter) if counter <= 26 else str(counter)
            return f'{n}.'
        if fmt == 'lowerLetter':
            n = chr(96 + counter) if counter <= 26 else str(counter)
            return f'{n})'

        # 罗马数字
        if fmt == 'upperRoman':
            return f'{self._to_roman(counter)}. '
        if fmt == 'lowerRoman':
            return f'{self._to_roman(counter).lower()}. '

        # 默认：阿拉伯数字
        if lvl_text and '%' in lvl_text:
            return lvl_text.replace('%1', str(counter)).replace('%2', str(counter))
        return f'{counter}. '

    @staticmethod
    def _to_roman(n):
        """数字转罗马数字"""
        val = [1000, 900, 500, 400, 100, 90, 50, 40, 10, 9, 5, 4, 1]
        sym = ['M', 'CM', 'D', 'CD', 'C', 'XC', 'L', 'XL', 'X', 'IX', 'V', 'IV', 'I']
        result = ''
        for i in range(len(val)):
            while n >= val[i]:
                result += sym[i]
                n -= val[i]
        return result

    def read_file(self, file_path):
        """读取文件，支持 .txt/.md/.docx 格式，.docx 会读取段落和表格内容"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        file_type = os.path.splitext(file_path)[1]
        if file_type not in FileHandler.type_map:
            raise ValueError(f"文件类型不支持: {file_type}")

        try:
            if file_type in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                return content, 'txt'

            elif file_type == '.docx':
                doc = Document(file_path)
                content_parts = []

                # 解析自动编号定义
                numbering_map = self._parse_numbering_formats(doc)
                list_counters = {}  # {numId: counter}

                # 1. 读取普通段落（含自动编号处理）
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue

                    # 检测是否是自动编号段落
                    num_pr = para._element.find('.//' + qn('w:numPr'))
                    if num_pr is not None:
                        num_id_el = num_pr.find(qn('w:numId'))
                        if num_id_el is not None:
                            num_id = num_id_el.get(qn('w:val'))

                            # 查找编号格式
                            fmt_info = numbering_map.get(num_id, {})
                            ilvl = '0'  # 默认0级
                            ilvl_el = num_pr.find(qn('w:ilvl'))
                            if ilvl_el is not None:
                                ilvl = ilvl_el.get(qn('w:val'))
                            level_info = fmt_info.get(ilvl, {})
                            fmt = level_info.get('fmt', 'decimal')
                            lvl_text = level_info.get('lvlText', '%1.')

                            # 项目符号不递增计数器
                            if fmt != 'bullet':
                                list_counters[num_id] = list_counters.get(num_id, 0) + 1
                            counter = list_counters.get(num_id, 1)

                            prefix = self._format_list_number(fmt, lvl_text, counter)
                            text = prefix + text

                    content_parts.append(text)

                # 2. 读取表格内容
                for table in doc.tables:
                    content_parts.append('--- 表格 ---')
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                row_text.append(cell_text)
                        if row_text:
                            content_parts.append(' | '.join(row_text))
                    content_parts.append('--- 表格结束 ---')

                content = '\n'.join(content_parts)
                return content, 'docx'

        except Exception as e:
            raise Exception(f"文件读取失败: {file_path}") from e

    def save_file(self, content, file_path, file_type):
        try:
            if file_type == 'txt':
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
            elif file_type == 'docx':
                doc = Document()
                # 设置默认样式的东西字体，避免中文字符显示为方框
                style = doc.styles['Normal']
                style_xml = style.element
                rPr = style_xml.find(qn('w:rPr'))
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    style_xml.append(rPr)
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = OxmlElement('w:rFonts')
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:eastAsia'), style.font.name or '宋体')

                for line in content.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)
                doc.save(file_path)
        except Exception as e:
            raise Exception(f"文件保存失败: {file_path}") from e






